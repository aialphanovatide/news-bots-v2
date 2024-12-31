# Function to extract text from a docx file

from docx import Document
import io
import re
from typing import Union
from werkzeug.datastructures import FileStorage

def extract_docx_content(docx_file: Union[FileStorage, io.BytesIO]) -> Union[str, None]:
    """
    Extract and clean text content from a DOCX file, including text from tables and paragraphs.

    Args:
        docx_file (Union[FileStorage, io.BytesIO]): The DOCX file to process, either as a FileStorage 
                                                   object (from form upload) or BytesIO object

    Returns:
        Union[str, None]: The cleaned extracted text if successful, None if unsuccessful
    """
    try:
        # Handle different input types
        if isinstance(docx_file, FileStorage):
            doc = Document(docx_file)
        else:
            doc = Document(docx_file)

        content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Only include non-empty paragraphs
                # Remove dots and ellipsis
                text = re.sub(r'\.{2,}', ' ', text)
                # Remove single dots that are not part of numbers
                text = re.sub(r'(?<!\d)\.(?!\d)', ' ', text)
                # Remove commas
                text = text.replace(',', '')
                content.append(text)

        # Extract text from tables
        for table in doc.tables:
            table_content = []
            
            # Process each row in the table
            for row in table.rows:
                # Extract and clean cell text
                row_text = [cell.text.strip() for cell in row.cells]
                # Only include rows that have non-empty cells
                if any(row_text):
                    table_content.append(' | '.join(row_text))
            
            # Add table content if any was found
            if table_content:
                content.append('\n'.join(table_content))

        # Check if any content was extracted
        if not content:
            return None

        # Join content and clean
        text = '\n\n'.join(content)

        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        # Remove spaces at start/end of lines
        text = re.sub(r'^\s+|\s+$', '', text, flags=re.MULTILINE)
        # Remove any null characters
        text = re.sub(r'\x00', '', text)
        # Remove any other control characters
        text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)

        return text if text.strip() else None

    except Exception:
        return None


# Example usage
# if __name__ == "__main__":
#     # Example usage
#     with open('app/services/news_creator/tools/ai_alpha_app.docx', 'rb') as docx_file:
#         result = extract_docx_content(docx_file)
#         print(result)

